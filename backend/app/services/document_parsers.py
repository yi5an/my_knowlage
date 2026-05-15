from dataclasses import dataclass, field
from pathlib import Path


class DocumentParseError(Exception):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    content_md: str
    content_text: str
    metadata: dict[str, str] = field(default_factory=dict)


class DocumentParser:
    supported_extensions: set[str]

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        raise NotImplementedError


class TextDocumentParser(DocumentParser):
    supported_extensions = {".txt"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        text = path.read_text(encoding="utf-8")
        return ParsedDocument(title=Path(original_name).stem, content_md=text, content_text=text)


class MarkdownDocumentParser(DocumentParser):
    supported_extensions = {".md", ".markdown"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        markdown = path.read_text(encoding="utf-8")
        title = _first_markdown_heading(markdown) or Path(original_name).stem
        return ParsedDocument(
            title=title,
            content_md=markdown,
            content_text=_strip_markdown(markdown),
        )


class CsvDocumentParser(DocumentParser):
    supported_extensions = {".csv"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        text = path.read_text(encoding="utf-8")
        return ParsedDocument(
            title=Path(original_name).stem,
            content_md=f"```csv\n{text}\n```",
            content_text=text,
            metadata={"table_format": "csv"},
        )


class PdfDocumentParser(DocumentParser):
    supported_extensions = {".pdf"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - dependency is installed in CI
            raise DocumentParseError("PDF parser dependency is not installed.") from exc

        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(page for page in pages if page.strip())
        if not text:
            raise DocumentParseError("No extractable text found in PDF.")
        return ParsedDocument(title=Path(original_name).stem, content_md=text, content_text=text)


class DocxDocumentParser(DocumentParser):
    supported_extensions = {".docx"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:  # pragma: no cover - dependency is installed in CI
            raise DocumentParseError("DOCX parser dependency is not installed.") from exc

        doc = DocxDocument(str(path))
        lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        text = "\n\n".join(lines)
        return ParsedDocument(title=Path(original_name).stem, content_md=text, content_text=text)


class XlsxDocumentParser(DocumentParser):
    supported_extensions = {".xlsx"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        try:
            from openpyxl import load_workbook  # type: ignore[import-untyped]
        except ImportError as exc:  # pragma: no cover - dependency is installed in CI
            raise DocumentParseError("XLSX parser dependency is not installed.") from exc

        workbook = load_workbook(path, read_only=True, data_only=True)
        sections: list[str] = []
        for sheet in workbook.worksheets:
            sections.append(f"## {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                sections.append("| " + " | ".join(values) + " |")
        text = "\n".join(sections)
        return ParsedDocument(
            title=Path(original_name).stem,
            content_md=text,
            content_text=text,
            metadata={"table_format": "xlsx"},
        )


class ImageOcrParser(DocumentParser):
    supported_extensions = {".png", ".jpg", ".jpeg", ".webp", ".tiff"}

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        del path, original_name
        raise DocumentParseError("Image OCR parser is defined but not implemented yet.")


class DocumentParserRouter:
    def __init__(self, parsers: list[DocumentParser] | None = None) -> None:
        self.parsers = parsers or [
            TextDocumentParser(),
            MarkdownDocumentParser(),
            CsvDocumentParser(),
            PdfDocumentParser(),
            DocxDocumentParser(),
            XlsxDocumentParser(),
            ImageOcrParser(),
        ]

    def parse(self, path: Path, original_name: str) -> ParsedDocument:
        extension = Path(original_name).suffix.lower()
        for parser in self.parsers:
            if extension in parser.supported_extensions:
                return parser.parse(path, original_name)
        raise DocumentParseError(f"Unsupported document type: {extension or 'unknown'}")


def _first_markdown_heading(markdown: str) -> str | None:
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip() or None
    return None


def _strip_markdown(markdown: str) -> str:
    lines = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            lines.append(stripped.lstrip("#").strip())
        else:
            lines.append(line.replace("**", "").replace("__", "").replace("`", ""))
    return "\n".join(lines)
